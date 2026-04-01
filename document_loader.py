from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document  # for LangChain Document objects
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
import os
import re

# --- Clean text function (for chunks) ---
def clean_text(text: str) -> str:
    text = text.replace('\t', ' ')             # Replace tabs with spaces
    text = re.sub(' +', ' ', text)             # Collapse multiple spaces
    return text.strip()

# --- Extract raw text from PDF ---
def extract_text_from_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text  # No cleaning here, keep original separators

# --- Extract raw text from DOCX ---
def extract_text_from_docx(file_path):
    doc = DocxDocument(file_path)
    return "\n".join([para.text for para in doc.paragraphs])  # No cleaning here either

# --- Load, split, and clean chunks ---
def load_and_split_document(file_path, chunk_size=1000, chunk_overlap=200):
    ext = os.path.splitext(file_path)[1].lower()

    # Extract raw text
    if ext == '.pdf':
        raw_text = extract_text_from_pdf(file_path)
    elif ext == '.docx':
        raw_text = extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format. Use .pdf or .docx")

    # Split text (preserving separators)
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ".", "!", "?", " ", ""]
    )
    chunks = splitter.create_documents([raw_text])

    # Clean individual chunks
    cleaned_chunks = [
        Document(page_content=clean_text(chunk.page_content), metadata=chunk.metadata)
        for chunk in chunks
    ]

    return cleaned_chunks
