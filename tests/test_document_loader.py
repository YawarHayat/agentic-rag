"""
Tests for document_loader.py

Uses temporary files to test real chunking logic without needing
pre-existing documents on disk.
"""

import pytest
import os
import tempfile
from langchain.schema import Document

import sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))


# ---------------------------------------------------------------------------
# clean_text
# ---------------------------------------------------------------------------

class TestCleanText:
    def test_collapses_multiple_spaces(self):
        from document_loader import clean_text
        assert clean_text("hello   world") == "hello world"

    def test_replaces_tabs_with_space(self):
        from document_loader import clean_text
        assert clean_text("hello\tworld") == "hello world"

    def test_strips_leading_trailing_whitespace(self):
        from document_loader import clean_text
        assert clean_text("  hello  ") == "hello"

    def test_empty_string(self):
        from document_loader import clean_text
        assert clean_text("") == ""

    def test_preserves_newlines(self):
        from document_loader import clean_text
        # clean_text only collapses spaces and tabs, not newlines
        result = clean_text("line one\nline two")
        assert "line one" in result
        assert "line two" in result


# ---------------------------------------------------------------------------
# load_and_split_document
# ---------------------------------------------------------------------------

class TestLoadAndSplitDocument:
    def test_raises_on_unsupported_format(self):
        from document_loader import load_and_split_document
        with pytest.raises(ValueError, match="Unsupported file format"):
            load_and_split_document("somefile.txt")

    def test_returns_list_of_documents_from_pdf(self, tmp_path):
        """Creates a minimal valid PDF and checks chunking output."""
        from document_loader import load_and_split_document

        # Write a minimal PDF with enough text to produce at least one chunk
        pdf_content = b"""%PDF-1.4
1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj
2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj
3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj
4 0 obj << /Length 44 >>
stream
BT /F1 12 Tf 100 700 Td (Hello World test document.) Tj ET
endstream
endobj
5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000274 00000 n
0000000368 00000 n
trailer << /Size 6 /Root 1 0 R >>
startxref
441
%%EOF"""
        pdf_file = tmp_path / "test.pdf"
        pdf_file.write_bytes(pdf_content)

        chunks = load_and_split_document(str(pdf_file))
        assert isinstance(chunks, list)
        # Every item should be a LangChain Document
        assert all(isinstance(c, Document) for c in chunks)

    def test_chunks_have_non_empty_content(self, tmp_path):
        from document_loader import load_and_split_document

        # Create a simple docx-like test using python-docx
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument()
            for i in range(20):
                doc.add_paragraph(
                    f"This is paragraph {i} with some meaningful content about the story. "
                    f"The detective Daisy investigated the mysterious case carefully."
                )
            docx_file = tmp_path / "test.docx"
            doc.save(str(docx_file))

            chunks = load_and_split_document(str(docx_file))
            assert len(chunks) >= 1
            assert all(len(c.page_content.strip()) > 0 for c in chunks)
        except ImportError:
            pytest.skip("python-docx not available")

    def test_chunk_size_respected(self, tmp_path):
        from document_loader import load_and_split_document

        try:
            from docx import Document as DocxDocument
            doc = DocxDocument()
            # Add enough text to force multiple chunks at chunk_size=200
            for i in range(50):
                doc.add_paragraph(
                    f"Sentence number {i}: The quick brown fox jumps over the lazy dog repeatedly. "
                )
            docx_file = tmp_path / "test_chunks.docx"
            doc.save(str(docx_file))

            chunks = load_and_split_document(str(docx_file), chunk_size=200, chunk_overlap=20)
            assert len(chunks) > 1
            # No single chunk should wildly exceed the chunk size
            for chunk in chunks:
                assert len(chunk.page_content) <= 400  # allow some overlap tolerance
        except ImportError:
            pytest.skip("python-docx not available")
